"""
Generate ERP Proposal Word document.
Run: python docs/proposals/generate_proposal.py
Output: docs/proposals/erp-proposal-2026-06-16.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour palette ──────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1e, 0x3a, 0x5f)
BLUE   = RGBColor(0x25, 0x63, 0xeb)
GREEN  = RGBColor(0x05, 0x96, 0x69)
LIGHT  = RGBColor(0xf0, 0xf4, 0xf8)
MUTED  = RGBColor(0x6b, 0x72, 0x80)
WHITE  = RGBColor(0xff, 0xff, 0xff)
YELLOW = RGBColor(0xff, 0xfb, 0xeb)
AMBER  = RGBColor(0x78, 0x35, 0x0f)
BLACK  = RGBColor(0x11, 0x18, 0x27)
LGRAY  = RGBColor(0xf3, 0xf4, 0xf6)
DGRAY  = RGBColor(0x37, 0x41, 0x51)

OUT_PATH = "docs/proposals/erp-proposal-2026-06-16.docx"


# ── Helpers ──────────────────────────────────────────────────────────────────

def rgb_hex(rgb: RGBColor) -> str:
    """Return 6-digit uppercase hex string from an RGBColor."""
    return str(rgb)  # RGBColor.__str__ returns e.g. '1E3A5F'


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(rgb))
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if side in kwargs:
            tag = OxmlElement(f'w:{side}')
            tag.set(qn('w:val'), kwargs[side].get('val', 'single'))
            tag.set(qn('w:sz'), str(kwargs[side].get('sz', 4)))
            tag.set(qn('w:space'), '0')
            tag.set(qn('w:color'), kwargs[side].get('color', 'auto'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def set_table_border(table, color='D1D5DB'):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{side}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), color)
        tblBorders.append(tag)
    tblPr.append(tblBorders)


def no_space_before(para):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)


def cell_para(cell, text='', bold=False, color=None, size=10,
              align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    para.alignment = align
    no_space_before(para)
    if text:
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
    return para


def add_cell_text(cell, text, bold=False, color=None, size=10,
                  align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    para = cell.add_paragraph()
    para.alignment = align
    no_space_before(para)
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return para


# ── Document setup ────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.font.color.rgb = BLACK


# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════

def cover_para(text, size=11, bold=False, color=None,
               align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    pPr = p._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(space_before))
    sp.set(qn('w:after'),  str(space_after))
    pPr.append(sp)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color or WHITE
    return p


cover_para('ERP', size=32, bold=True, space_before=120, space_after=20)
cover_para('Enterprise Resource Planning System', size=22, bold=True, space_after=10)
cover_para('Software Proposal & Module Selection', size=13, space_after=14)
cover_para(
    'Philippine-Compliant  ·  BIR-Ready  ·  Multi-Branch  ·  Web-Based  ·  TRAIN Act'
    '  ·  SSS · PhilHealth · Pag-IBIG',
    size=10, color=RGBColor(0xb0, 0xc4, 0xde), space_after=60
)

# Meta table (2 col × 4 row)
meta = doc.add_table(rows=4, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta.style = 'Table Grid'
set_table_border(meta, color='FFFFFF')

meta_data = [
    ('PREPARED FOR',   ''),
    ('PROPOSAL DATE',  'June 16, 2026'),
    ('PREPARED BY',    ''),
    ('VALID UNTIL',    ''),
]
for i, (label, value) in enumerate(meta_data):
    for j, (txt, bld, sz, clr) in enumerate([
        (label, False, 8,  RGBColor(0xb0, 0xc4, 0xde)),
        (value, True,  11, WHITE),
    ]):
        cell = meta.cell(i, j)
        set_cell_bg(cell, NAVY)
        p = cell_para(cell, txt, bold=bld, color=clr, size=sz)
        if not value and j == 1:
            # blank field — draw underline
            run = cell.paragraphs[0].add_run('_' * 30)
            run.font.color.rgb = RGBColor(0x60, 0x80, 0xa0)
            run.font.size = Pt(11)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# HELPERS FOR BODY
# ════════════════════════════════════════════════════════════════════════════

def section_heading(number, title):
    """Navy numbered section header."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '160')
    sp.set(qn('w:after'),  '60')
    pPr.append(sp)

    # Number circle via square-bracketed bold
    r1 = p.add_run(f'  {number}  ')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = WHITE
    # Shade the number run navy — use highlight workaround via character shading
    rPr = r1._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(NAVY))
    rPr.append(shd)

    r2 = p.add_run(f'   {title}')
    r2.bold = True
    r2.font.size = Pt(14)
    r2.font.color.rgb = NAVY

    # Bottom border on paragraph
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1E3A5F')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def body_para(text, size=10, bold=False, color=None,
              space_before=0, space_after=80, italic=False):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(space_before))
    sp.set(qn('w:after'),  str(space_after))
    pPr.append(sp)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color or DGRAY
    return p


def info_box(text):
    """Blue left-border info paragraph (table workaround)."""
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Left thin accent column
    accent = t.cell(0, 0)
    accent.width = Cm(0.35)
    set_cell_bg(accent, BLUE)
    cell_para(accent, '')
    # Content column
    content = t.cell(0, 1)
    set_cell_bg(content, LIGHT)
    p = cell_para(content, text, size=10, color=NAVY)
    pPr = p._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '60'); sp.set(qn('w:after'), '60')
    pPr.append(sp)
    # Remove outer table border
    set_table_border(t, color='F0F4F8')
    doc.add_paragraph()  # spacer


def note_box(text):
    """Amber note box."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0, 0)
    set_cell_bg(cell, YELLOW)
    set_table_border(t, color='FCD34D')
    p = cell_para(cell, text, size=9.5, color=AMBER)
    pPr = p._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '60'); sp.set(qn('w:after'), '60')
    pPr.append(sp)
    doc.add_paragraph()


def plain_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_border(t)
    # Header row
    for j, h in enumerate(headers):
        cell = t.cell(0, j)
        set_cell_bg(cell, LIGHT)
        p = cell_para(cell, h, bold=True, color=NAVY, size=9.5)
        pPr = p._p.get_or_add_pPr()
        sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '40'); sp.set(qn('w:after'), '40')
        pPr.append(sp)
    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.cell(i + 1, j)
            if i % 2 == 1:
                set_cell_bg(cell, LGRAY)
            bold = (j == 0)
            p = cell_para(cell, val, bold=bold, color=DGRAY, size=9.5)
            pPr = p._p.get_or_add_pPr()
            sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '40'); sp.set(qn('w:after'), '40')
            pPr.append(sp)
    if col_widths:
        for i, row_cells in enumerate(t.rows):
            for j, cell in enumerate(row_cells.cells):
                if j < len(col_widths):
                    cell.width = Cm(col_widths[j])
    doc.add_paragraph()
    return t


# ════════════════════════════════════════════════════════════════════════════
# MODULE GROUP BUILDER
# ════════════════════════════════════════════════════════════════════════════

def module_group(title, icon, modules, is_foundation=False):
    """
    modules = list of (name, description, features_list, is_checked_or_included)
    """
    bg = GREEN if is_foundation else NAVY

    # Group header table
    hdr = doc.add_table(rows=1, cols=1)
    set_table_border(hdr, color='FFFFFF')
    hcell = hdr.cell(0, 0)
    set_cell_bg(hcell, bg)
    badge = '  ✓ INCLUDED IN ALL PACKAGES' if is_foundation else ''
    p = cell_para(hcell, f'{icon}  {title}{badge}',
                  bold=True, color=WHITE, size=10)
    pPr = p._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '60'); sp.set(qn('w:after'), '60')
    pPr.append(sp)

    # Module rows table — cols: checkbox | name+desc+features | cost
    n_rows = 1 + len(modules)  # header + data
    t = doc.add_table(rows=n_rows, cols=3)
    set_table_border(t)

    # Column header row
    col_hdrs = ['', 'Module', 'Cost (PHP)']
    col_widths_cm = [1.2, 12.5, 2.8]
    for j, h in enumerate(col_hdrs):
        cell = t.cell(0, j)
        set_cell_bg(cell, LGRAY)
        p = cell_para(cell, h, bold=True, color=MUTED, size=8.5)
        pPr = p._p.get_or_add_pPr()
        sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '30'); sp.set(qn('w:after'), '30')
        pPr.append(sp)

    # Data rows
    for i, (name, desc, features, included) in enumerate(modules):
        row_bg = RGBColor(0xf0, 0xfd, 0xf4) if is_foundation else None

        # Checkbox / tick cell
        chk_cell = t.cell(i + 1, 0)
        if row_bg: set_cell_bg(chk_cell, row_bg)
        chk_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        mark = '✓' if is_foundation else '☐'
        clr  = GREEN if is_foundation else BLACK
        p = cell_para(chk_cell, mark, bold=True, color=clr, size=12,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        pPr = p._p.get_or_add_pPr()
        sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '60'); sp.set(qn('w:after'), '0')
        pPr.append(sp)

        # Name + desc + features cell
        content_cell = t.cell(i + 1, 1)
        if row_bg: set_cell_bg(content_cell, row_bg)
        content_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Module name
        p_name = content_cell.paragraphs[0]
        p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
        no_space_before(p_name)
        pPr = p_name._p.get_or_add_pPr()
        sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '60'); sp.set(qn('w:after'), '20')
        pPr.append(sp)
        r = p_name.add_run(name)
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = BLACK

        # Description
        p_desc = content_cell.add_paragraph()
        no_space_before(p_desc)
        pPr = p_desc._p.get_or_add_pPr()
        sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '30')
        pPr.append(sp)
        r2 = p_desc.add_run(desc)
        r2.italic = True; r2.font.size = Pt(9); r2.font.color.rgb = MUTED

        # Feature tags (comma-separated on one line)
        if features:
            p_feat = content_cell.add_paragraph()
            no_space_before(p_feat)
            pPr = p_feat._p.get_or_add_pPr()
            sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '60')
            pPr.append(sp)
            tags_text = '  ·  '.join(features)
            r3 = p_feat.add_run(tags_text)
            r3.font.size = Pt(8.5); r3.font.color.rgb = RGBColor(0x4b, 0x55, 0x63)

        # Cost cell
        cost_cell = t.cell(i + 1, 2)
        if row_bg: set_cell_bg(cost_cell, row_bg)
        cost_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        cost_txt = 'Included' if is_foundation else ''
        p_cost = cell_para(cost_cell, cost_txt, color=MUTED, size=9,
                           align=WD_ALIGN_PARAGRAPH.RIGHT)
        pPr = p_cost._p.get_or_add_pPr()
        sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '60'); sp.set(qn('w:after'), '0')
        pPr.append(sp)

    # Set column widths
    for row in t.rows:
        for j, cell in enumerate(row.cells):
            cell.width = Cm(col_widths_cm[j])

    doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════
# SECTION 1 — EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════════

section_heading('1', 'Executive Summary')

info_box(
    'This proposal outlines a fully custom-built Enterprise Resource Planning (ERP) system '
    'tailored for Philippine small and medium-sized enterprises. The system consolidates all '
    'core business functions — accounting, purchasing, sales, inventory, payroll, and production '
    '— into a single, integrated, web-based platform accessible from any device.'
)

body_para(
    'The ERP is built around Philippine regulatory requirements from the ground up: BIR VAT and '
    'withholding tax rules, TRAIN Act income tax tables, and mandatory government contributions '
    '(SSS, PhilHealth, Pag-IBIG). All financial reports and data exports are formatted for BIR compliance.'
)
body_para(
    'The system is modular. A core Foundation package is always included. Every additional module '
    'is optional — the client selects only what the business needs today, and more modules can be '
    'added at any time without rebuilding the system.'
)


# ════════════════════════════════════════════════════════════════════════════
# SECTION 2 — SYSTEM OVERVIEW
# ════════════════════════════════════════════════════════════════════════════

section_heading('2', 'System Overview')

body_para(
    'All modules share a common foundation: one database, unified user access control, a full audit '
    'trail, and a consistent interface across every screen. The table below summarizes the functional '
    'areas available.'
)

plain_table(
    headers=['Functional Area', 'What It Covers'],
    rows=[
        ('Foundation',
         'Chart of accounts, double-entry journals, accounting periods, user roles, multi-branch, audit trail'),
        ('Financial Management',
         'Accounts receivable, accounts payable, cash receipts, disbursements, fixed assets, bank reconciliation, budgeting'),
        ('BIR Tax Compliance',
         'VAT input/output, withholding tax codes, BIR summary lists, alphalist, form generation'),
        ('Sales Cycle',
         'Customer master, sales quotations, sales orders, delivery receipts, price lists'),
        ('Procurement Cycle',
         'Vendor master, purchase requests, purchase orders, receiving reports, 3-way matching'),
        ('Inventory Management',
         'Item master, stock movements, inventory valuation, physical count, reorder alerts'),
        ('Human Resources & Payroll',
         'Employee records, payroll computation, government remittances, leave management, payslips'),
        ('Production / Manufacturing',
         'Bill of materials, work orders, production costing, variance reports'),
        ('Reporting & Analytics',
         'Financial statements, aging reports, cash flow, management dashboards, Excel/CSV export'),
        ('System & Integrations',
         'Document attachments, email alerts, bulk import, multi-currency, API access'),
    ],
    col_widths=[5.5, 11.0],
)


# ════════════════════════════════════════════════════════════════════════════
# SECTION 3 — MODULE SELECTION
# ════════════════════════════════════════════════════════════════════════════

section_heading('3', 'Module Selection')

body_para(
    'Please place a check mark next to every module you wish to include. The Foundation block is '
    'mandatory and included in all packages. All other modules are optional add-ons.'
)

note_box(
    'Instructions: Check the modules your business needs. If you are unsure about a module, leave '
    'it unchecked — we can discuss it separately. Additional modules may be added in a future phase '
    'at any time. Pricing and timeline will be finalized based on your selections.'
)

# ── Foundation ──
module_group('Foundation — Always Included', '🏛', [
    ('Chart of Accounts',
     'Hierarchical general ledger accounts: Asset, Liability, Equity, Revenue, Expense',
     ['Account codes & names', 'Parent-child hierarchy', 'Account types & classification'],
     True),
    ('Journal Entries & General Ledger',
     'Double-entry bookkeeping, manual adjustments, reversals, and auto-generated entries from all transactions',
     ['Draft → Posted workflow', 'Auto-balance validation', 'Reversing entries', 'Adjustment / closing / reclassification types'],
     True),
    ('Accounting Periods',
     'Monthly open/closed fiscal period management; no transactions can post to a closed period',
     ['Open / Close periods', 'Year-end close', 'Period-lock enforcement'],
     True),
    ('User Management & Role-Based Access Control',
     'Multi-user system with roles controlling what each person can view or change',
     ['Admin / Accountant / Staff / Viewer roles', 'Branch-level access', 'Account lockout after failed logins'],
     True),
    ('Multi-Branch Support',
     'Multiple business locations under one system, each with its own data scope and access permissions',
     ['Branch picker per session', 'Branch-scoped transactions', 'Cross-branch reporting'],
     True),
    ('Audit Trail & Approval Workflow',
     'Every action is permanently recorded. Sensitive changes require approval before taking effect.',
     ['Full audit log (who, when, what changed)', 'Pending approval queue', 'Self-approval prevention'],
     True),
    ('Dashboard & Company Settings',
     'Home screen with key metrics and pending action items; company profile and system settings',
     ['Pending approvals badge', 'Company name & logo', 'System notifications'],
     True),
], is_foundation=True)

# ── Financial Management ──
module_group('Financial Management', '💰', [
    ('Accounts Receivable — Sales Invoices',
     'Issue official sales invoices to customers; track outstanding balances; auto-post journal entries',
     ['Invoice numbering (SI-YYYY-MM-####)', 'VAT-inclusive line items', 'WT receivable per line', 'Draft → Posted → Void', 'PDF print'],
     False),
    ('Accounts Payable — Vendor Bills',
     'Record vendor invoices; track outstanding payables; auto-post journal entries with VAT and withholding tax',
     ['Bill numbering (AP-YYYY-MM-####)', 'VAT input tax per line', 'Withholding tax per line', 'Draft → Posted → Void', 'Attachment support'],
     False),
    ('Cash Disbursement Vouchers',
     'Record cash or check payments to vendors; supports full or partial payment against AP bills or direct expenses',
     ['CDV numbering', 'Apply to AP bills (partial / full)', 'Direct expense lines', 'Check number tracking', 'Auto journal entry'],
     False),
    ('Cash Receipts & Collections',
     'Record customer payments via cash, check, or bank transfer; apply collections to outstanding invoices',
     ['Collection receipt numbering', 'Cash / check / bank transfer', 'Apply to SI invoices', 'Auto journal entry'],
     False),
    ('Fixed Asset Management',
     'Maintain a register of all company assets; automatically compute and post monthly depreciation',
     ['Asset register (code, description, cost, date)', 'Straight-line & declining balance', 'Auto depreciation journal entries', 'Disposal / write-off workflow'],
     False),
    ('Bank Reconciliation',
     'Match book entries against the bank statement; identify outstanding checks and deposits in transit',
     ['Bank statement import (CSV)', 'Match / unmatch entries', 'Outstanding check list', 'Reconciliation report'],
     False),
    ('Budgeting & Cost Control',
     'Set annual budgets per account; compare budget vs. actual with variance analysis',
     ['Annual budget entry per account / month', 'Budget vs. Actual report', 'Over-budget alerts', 'Department / cost center tagging'],
     False),
])

# ── BIR Compliance ──
module_group('BIR Tax Compliance', '🧾', [
    ('VAT Management',
     'Manage VAT categories (12%, 0%, exempt, non-VAT); auto-extract VAT from line amounts; BIR-compliant treatment of input and output tax',
     ['VAT category master', 'VAT-inclusive extraction', 'Input vs. output VAT accounts', 'Summary List of Sales (SLS)', 'Summary List of Purchases (SLP)'],
     False),
    ('Withholding Tax Management',
     'Expanded withholding tax per BIR ATC codes; auto-compute on purchases and payments; track payables to BIR',
     ['WT code master (ATC codes)', 'Per-line WT computation', 'WT receivable & payable tracking', 'Alphalist of Payees (BIR 1604-E)', 'BIR 2307 certificate tracking'],
     False),
    ('BIR Form Generation',
     'Generate data files and summaries ready for eBIRForms or EFPS submission',
     ['VAT Return summary (BIR 2550M / 2550Q)', 'Expanded WT summary (1601-EQ)', 'Alphalist data file (1604-CF)', 'CSV / Excel export for upload'],
     False),
])

# ── Sales Cycle ──
module_group('Customer & Sales Cycle', '📦', [
    ('Customer Master',
     'Centralized customer database shared across all branches',
     ['Customer code, name, TIN, contact', 'Default payment terms', 'Default VAT category', 'Customer account statement'],
     False),
    ('Sales Quotation',
     'Prepare and send price quotations to customers; convert approved quotations to Sales Orders',
     ['Quotation numbering', 'Line items with pricing', 'Validity date', 'Convert to Sales Order', 'PDF generation'],
     False),
    ('Sales Order',
     'Record confirmed customer orders; track fulfillment status; trigger Delivery Receipts',
     ['SO numbering (SO-YYYY-MM-####)', 'Open / Partial / Fulfilled status', 'Backorder management', 'Link to Delivery Receipts'],
     False),
    ('Delivery Receipt',
     'Record goods or services delivered to customers; reduce inventory stock; trigger billing',
     ['DR numbering (DR-YYYY-MM-####)', 'Partial delivery support', 'Inventory out movement', 'Link to Sales Order & Sales Invoice', 'PDF delivery slip'],
     False),
    ('Price Lists',
     'Manage multiple price tiers per item; automatically apply the correct price on sales documents',
     ['Price list by customer tier or group', 'Effective date management', 'Discount rules'],
     False),
])

# ── Procurement ──
module_group('Vendor & Procurement Cycle', '🛒', [
    ('Vendor Master',
     'Centralized supplier database with BIR-required information and default WT codes',
     ['Vendor code, name, TIN, contact', 'Default payment terms', 'Default WT codes', 'Vendor account statement'],
     False),
    ('Purchase Request',
     'Internal requisition for goods or services; routed for approval before any purchasing action',
     ['PR numbering (PR-YYYY-MM-####)', 'Requestor & department', 'Approval workflow', 'Convert to Purchase Order'],
     False),
    ('Purchase Order',
     'Formal order issued to a vendor; tracks ordered vs. received quantities',
     ['PO numbering (PO-YYYY-MM-####)', 'Open / Partial / Fulfilled status', 'PDF for vendor', 'Link to Receiving Report'],
     False),
    ('Receiving Report',
     'Record goods received from the vendor; increase inventory stock; trigger AP billing',
     ['RR numbering (RR-YYYY-MM-####)', 'Partial receipt support', 'Inventory in movement', 'Link to PO & AP Bill'],
     False),
    ('3-Way Matching',
     'Internal control: verify that the Purchase Order, Receiving Report, and AP invoice all agree before authorizing payment',
     ['PO vs. RR vs. AP Bill comparison', 'Discrepancy alerts', 'Block payment on mismatch'],
     False),
])

# ── Inventory ──
module_group('Inventory Management', '📊', [
    ('Item / Product Master',
     'Central catalog of all goods and services the business buys and sells',
     ['Item code, name, unit of measure', 'Cost method (FIFO / Weighted Average)', 'Reorder point & minimum stock level', 'Sales price & purchase price'],
     False),
    ('Stock Movements & Ledger',
     'Every transaction creates a stock movement: receipt, delivery, transfer, adjustment. Running balance maintained per item per branch.',
     ['Stock in / out / transfer / adjustment', 'Branch-level tracking', 'Running stock balance', 'FIFO / Weighted Average valuation'],
     False),
    ('Physical Count & Stock Take',
     'Conduct periodic inventory counts; compare physical count vs. book balance; post adjustment journal entries',
     ['Count sheet generation', 'Variance report (book vs. actual)', 'Auto inventory adjustment JE'],
     False),
    ('Inventory Reports',
     'Inventory aging, valuation, movement history, and low-stock alert reports',
     ['Stock balance by item / branch', 'Inventory valuation report', 'Slow-moving / fast-moving analysis', 'Reorder alert list'],
     False),
])

# ── HR & Payroll ──
module_group('Human Resources & Payroll', '👥', [
    ('Employee Master',
     'Complete employee records including all government-mandated identification and employment details',
     ['Employee code, name, position, department', 'SSS, PhilHealth, Pag-IBIG numbers', 'TIN & BIR tax status (S / ME / etc.)', 'Employment type (regular / probationary / contractual)', 'Bank account for payroll crediting'],
     False),
    ('Payroll Computation',
     'Semi-monthly or monthly payroll processing; TRAIN Act-compliant income tax tables; all mandated deductions auto-computed',
     ['Basic pay, allowances, overtime, holiday pay', 'SSS, PhilHealth, Pag-IBIG deductions', 'Withholding tax on compensation (TRAIN table)', 'De minimis benefits', 'Net pay computation', 'Payslip generation (PDF)'],
     False),
    ('Government Remittances',
     'Generate all required contribution reports and schedules for government agencies',
     ['SSS R-3 report', 'PhilHealth RF-1 report', 'Pag-IBIG remittance list', 'BIR 1601-C (WT on compensation)', 'BIR 2316 (employee tax certificate)', 'Alphalist of employees (1604-CF)'],
     False),
    ('13th Month Pay',
     'Automatically compute 13th month pay per employee based on basic pay earned; pro-rated for new hires and resignees',
     ['Annual 13th month computation', 'Pro-rated for partial year', 'BIR exemption threshold tracking'],
     False),
    ('Leave Management',
     'Track employee leave balances and applications; deduct leave without pay from payroll automatically',
     ['Leave type setup (VL, SL, EL, Maternity, Paternity)', 'Leave balance tracking', 'Leave application & approval', 'Leave without pay deduction to payroll'],
     False),
])

# ── Production ──
module_group('Production / Manufacturing', '🏭', [
    ('Bill of Materials (BOM)',
     'Define the raw materials and components required to produce each finished product',
     ['Multi-level BOM', 'Component quantities per unit', 'Standard cost computation'],
     False),
    ('Work Orders',
     'Authorize and track production runs; consume raw materials from inventory; record finished goods produced',
     ['Work order numbering', 'Material issuance from inventory', 'Finished goods receipt', 'Scrap / wastage tracking'],
     False),
    ('Production Costing & Variance Reports',
     'Compare standard vs. actual production cost; auto-generate WIP and COGS journal entries',
     ['Standard vs. actual cost comparison', 'Material & usage variance', 'WIP → Finished Goods journal entries', 'Production cost report'],
     False),
])

# ── Reporting ──
module_group('Reporting & Analytics', '📈', [
    ('Financial Statements',
     'Trial Balance, Income Statement, Balance Sheet, and Statement of Cash Flows — per period, per branch, or consolidated',
     ['Trial Balance', 'Income Statement (P&L)', 'Balance Sheet', 'Statement of Cash Flows', 'Excel / PDF export'],
     False),
    ('AR / AP Aging Reports',
     'Outstanding receivables and payables aged by 30 / 60 / 90 / 90+ day buckets',
     ['AR Aging by customer', 'AP Aging by vendor', 'As-of-date aging', 'Excel export'],
     False),
    ('Management Reports',
     'Operational reports for management decision-making beyond standard financial statements',
     ['Sales summary by period / branch', 'Purchase summary', 'Top customers & vendors', 'Cash position report'],
     False),
])

# ── System & Integrations ──
module_group('System Features & Integrations', '⚙', [
    ('Document Attachments',
     'Attach scanned receipts, contracts, or photos directly to any transaction record',
     ['File upload per transaction', 'PDF, JPG, PNG support', 'View / download in-app'],
     False),
    ('Email Notifications',
     'Automated email alerts for pending approvals, overdue balances, and key system events',
     ['Approval request emails', 'Overdue AR / AP alerts', 'Low stock alerts', 'Payroll processed notification'],
     False),
    ('Bulk Data Import',
     'Import master data and opening balances from Excel templates to accelerate system setup',
     ['Customer / vendor import', 'Chart of accounts import', 'Opening balance upload', 'Inventory count import'],
     False),
    ('Multi-Currency Support',
     'Transact in USD, EUR, and other foreign currencies; auto-post foreign exchange gain or loss on settlement',
     ['Currency master with exchange rates', 'BSP rate management', 'Forex gain / loss journal entries', 'Multi-currency AR / AP'],
     False),
    ('API Access & Integration',
     'REST API endpoints for connecting the ERP to POS systems, e-commerce platforms, or third-party tools',
     ['Authenticated REST endpoints', 'Read / write access per module', 'Webhook support'],
     False),
])


# ════════════════════════════════════════════════════════════════════════════
# SECTION 4 — INVESTMENT SUMMARY
# ════════════════════════════════════════════════════════════════════════════

section_heading('4', 'Investment Summary')

body_para(
    'The investment below will be completed based on the modules checked in Section 3. '
    'Pricing is presented as a one-time development fee. An optional annual support and '
    'maintenance agreement is available separately.'
)

sum_rows = [
    ('Foundation Package (always included)',             ''),
    ('Financial Management — selected modules',          ''),
    ('BIR Tax Compliance — selected modules',            ''),
    ('Sales & Procurement Cycle — selected modules',     ''),
    ('Inventory Management — selected modules',          ''),
    ('Human Resources & Payroll — selected modules',     ''),
    ('Production / Manufacturing — selected modules',    ''),
    ('Reporting, System Features & Integrations',        ''),
    ('Data Migration & System Setup',                    ''),
    ('User Training & Go-Live Support',                  ''),
    ('Annual Support & Maintenance (optional)',           ''),
]

t = doc.add_table(rows=1 + len(sum_rows) + 1, cols=2)
set_table_border(t)
t.alignment = WD_TABLE_ALIGNMENT.LEFT

# Header
for j, h in enumerate(['Item', 'Amount (PHP)']):
    cell = t.cell(0, j)
    set_cell_bg(cell, LIGHT)
    p = cell_para(cell, h, bold=True, color=NAVY, size=9.5,
                  align=WD_ALIGN_PARAGRAPH.RIGHT if j else WD_ALIGN_PARAGRAPH.LEFT)
    pPr = p._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '40'); sp.set(qn('w:after'), '40')
    pPr.append(sp)

# Data rows
for i, (label, amount) in enumerate(sum_rows):
    row_idx = i + 1
    is_optional = 'optional' in label.lower()
    cell0 = t.cell(row_idx, 0)
    cell1 = t.cell(row_idx, 1)
    if i % 2 == 1:
        set_cell_bg(cell0, LGRAY); set_cell_bg(cell1, LGRAY)
    clr = MUTED if is_optional else DGRAY
    sz  = 9 if is_optional else 9.5
    p0 = cell_para(cell0, label, color=clr, size=sz, italic=is_optional)
    p1 = cell_para(cell1, '___________', color=clr, size=sz,
                   align=WD_ALIGN_PARAGRAPH.RIGHT, italic=is_optional)
    for p in [p0, p1]:
        pPr = p._p.get_or_add_pPr()
        sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '40'); sp.set(qn('w:after'), '40')
        pPr.append(sp)

# Total row
total_row = len(sum_rows) + 1
for j, txt in enumerate(['TOTAL PROJECT INVESTMENT', 'PHP ___________']):
    cell = t.cell(total_row, j)
    set_cell_bg(cell, NAVY)
    align = WD_ALIGN_PARAGRAPH.RIGHT if j else WD_ALIGN_PARAGRAPH.LEFT
    p = cell_para(cell, txt, bold=True, color=WHITE, size=11, align=align)
    pPr = p._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '60'); sp.set(qn('w:after'), '60')
    pPr.append(sp)

# Column widths
for row in t.rows:
    row.cells[0].width = Cm(12.5)
    row.cells[1].width = Cm(4.0)

doc.add_paragraph()
note_box(
    'Payment Terms (typical arrangement): 50% upon contract signing · '
    '30% upon successful User Acceptance Testing (UAT) · 20% upon go-live. '
    'Final terms to be agreed in the contract.'
)


# ════════════════════════════════════════════════════════════════════════════
# SECTION 5 — PROJECT TIMELINE
# ════════════════════════════════════════════════════════════════════════════

section_heading('5', 'Estimated Project Timeline')

body_para(
    'The timeline below reflects the full-scope scenario. Phases for modules not selected will be '
    'skipped, reducing the total duration proportionally. A core accounting + AR/AP + Sales + '
    'Procurement package (without HR or Production) is typically completed in 20–24 weeks.'
)

phases = [
    ('Phase 1 — Foundation & Core Accounting', 'Weeks 1 – 4',
     'Chart of accounts, journal entries, user accounts and roles, branch setup, accounting period management, audit trail, dashboard'),
    ('Phase 2 — Financial Management & BIR Compliance', 'Weeks 5 – 10',
     'Accounts receivable, accounts payable, receipts, cash disbursements, VAT management, withholding tax, BIR reports, fixed assets, bank reconciliation'),
    ('Phase 3 — Sales & Procurement Cycles', 'Weeks 11 – 16',
     'Customer and vendor master, sales quotations, sales orders, delivery receipts, purchase requests, purchase orders, receiving reports, 3-way matching, price lists'),
    ('Phase 4 — Inventory Management', 'Weeks 17 – 20',
     'Item master, stock movements and ledger, inventory valuation (FIFO / Weighted Average), physical count, reorder alerts, inventory reports'),
    ('Phase 5 — Human Resources & Payroll', 'Weeks 21 – 27',
     'Employee master, payroll computation (TRAIN Act), government remittances (SSS / PhilHealth / Pag-IBIG), 13th month pay, leave management, payslip generation'),
    ('Phase 6 — Production & Advanced Features', 'Weeks 28 – 33',
     'Bill of materials, work orders, production costing, variance reports, budgeting, advanced reporting, bulk data import, email notifications, API access, multi-currency'),
    ('Phase 7 — User Acceptance Testing & Go-Live', 'Weeks 34 – 37',
     'Full UAT with key users per module, bug resolution, historical data migration, user training, go-live support, and post-launch monitoring'),
]

t = doc.add_table(rows=len(phases), cols=2)
set_table_border(t, color='E5E7EB')
t.alignment = WD_TABLE_ALIGNMENT.LEFT

for i, (phase, duration, desc) in enumerate(phases):
    # Left accent column
    dot_cell = t.cell(i, 0)
    dot_cell.width = Cm(0.6)
    set_cell_bg(dot_cell, BLUE if i % 2 == 0 else NAVY)
    cell_para(dot_cell, '', size=8)

    # Content column
    content = t.cell(i, 1)
    if i % 2 == 1:
        set_cell_bg(content, LGRAY)

    p1 = content.paragraphs[0]
    no_space_before(p1)
    pPr = p1._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '60'); sp.set(qn('w:after'), '20')
    pPr.append(sp)
    r = p1.add_run(phase)
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = NAVY

    p2 = content.add_paragraph()
    no_space_before(p2)
    pPr = p2._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '20')
    pPr.append(sp)
    r2 = p2.add_run(duration)
    r2.bold = True; r2.font.size = Pt(9); r2.font.color.rgb = BLUE

    p3 = content.add_paragraph()
    no_space_before(p3)
    pPr = p3._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '60')
    pPr.append(sp)
    r3 = p3.add_run(desc)
    r3.font.size = Pt(9); r3.font.color.rgb = DGRAY

for row in t.rows:
    row.cells[0].width = Cm(0.5)
    row.cells[1].width = Cm(16.0)

doc.add_paragraph()
note_box(
    'Note: All timelines are estimates. Final schedule will be confirmed in the project contract '
    'based on selected modules, client resource availability, and data migration complexity.'
)


# ════════════════════════════════════════════════════════════════════════════
# SECTION 6 — TECHNICAL SPECIFICATIONS
# ════════════════════════════════════════════════════════════════════════════

section_heading('6', 'Technical Specifications')

plain_table(
    headers=['Specification', 'Details'],
    rows=[
        ('Platform',        'Web-based application — no software installation required on user devices'),
        ('Device Access',   'Desktop computer, laptop, tablet, and mobile phone (fully responsive interface)'),
        ('Hosting Options', 'Cloud-hosted (managed server) or on-premise installation — client\'s choice'),
        ('Database',        'SQLite for standard SME scale · PostgreSQL available for high-volume or enterprise needs'),
        ('Security',        'HTTPS encryption, CSRF protection, role-based access control, account lockout, full audit trail'),
        ('BIR Compliance',  'CAS-compliant design, TRAIN Act income tax tables, BIR-format reports and data files for eBIRForms / EFPS'),
        ('Philippine Labor','DOLE-aligned leave types, SSS / PhilHealth / Pag-IBIG contribution tables updated per current regulations'),
        ('Data Backup',     'Automated daily backups; backup file downloadable by the system administrator at any time'),
        ('Browser Support', 'Google Chrome, Microsoft Edge, Mozilla Firefox, Apple Safari (latest versions)'),
        ('Data Exports',    'All reports exportable to Microsoft Excel (.xlsx) and CSV; select documents generate PDF'),
    ],
    col_widths=[5.0, 11.5],
)


# ════════════════════════════════════════════════════════════════════════════
# SECTION 7 — TERMS & CONDITIONS
# ════════════════════════════════════════════════════════════════════════════

section_heading('7', 'Terms & Conditions')

terms = [
    ('Scope',
     'The scope of this engagement is limited to the modules checked in Section 3. Any module or feature '
     'not checked at signing requires a separate change order, which may adjust the price and timeline.'),
    ('Project Coordinator',
     'The client will designate a project coordinator who will be available for requirements clarification, '
     'review of deliverables, UAT participation, and sign-off at each phase milestone.'),
    ('Training',
     'Training will be provided for up to ______ user accounts. Additional training sessions beyond those '
     'included are available at a separate rate.'),
    ('Data Migration',
     'Data migration from existing systems (spreadsheets, other software) is included for the agreed data scope. '
     'Data preparation, cleaning, and verification before migration is the responsibility of the client.'),
    ('Support & Maintenance',
     'Annual support and maintenance covers bug fixes and minor UI enhancements. Development of new modules or '
     'significant new features is scoped and billed separately.'),
    ('Intellectual Property',
     'The developer retains ownership of the underlying platform and source code. The client receives a '
     'perpetual, non-transferable license to operate the deployed system for their own business use.'),
    ('Proposal Validity',
     'This proposal is valid for 30 days from the proposal date above. Pricing and timelines are subject '
     'to revision after this period.'),
    ('Governing Law',
     'Philippine law governs this agreement. Any disputes shall first be resolved through good-faith '
     'negotiation, then mediation, before any formal proceedings.'),
]

for i, (title, text) in enumerate(terms, 1):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '40')
    sp.set(qn('w:after'), '40')
    pPr.append(sp)
    r1 = p.add_run(f'{i}.  {title}: ')
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = NAVY
    r2 = p.add_run(text)
    r2.font.size = Pt(10); r2.font.color.rgb = DGRAY


# ════════════════════════════════════════════════════════════════════════════
# SECTION 8 — ACCEPTANCE & SIGNATURE
# ════════════════════════════════════════════════════════════════════════════

section_heading('8', 'Acceptance & Authorization')

info_box(
    'By signing below, the Client confirms that the modules checked in Section 3 represent the agreed '
    'scope of work and authorizes the Developer to proceed on the terms stated in this proposal. '
    'Both parties agree that this document, once signed, forms the basis of the project contract.'
)

sig = doc.add_table(rows=1, cols=2)
set_table_border(sig, color='FFFFFF')
sig.alignment = WD_TABLE_ALIGNMENT.LEFT

for j, (role, party) in enumerate([
    ('Client — Authorized Representative', 'CLIENT'),
    ('Developer — Authorized Representative', 'DEVELOPER'),
]):
    cell = sig.cell(0, j)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    def sp_line(cell, txt, bold=False, color=None, size=9.5, before=40, after=40):
        p = cell.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:before'), str(before))
        sp.set(qn('w:after'), str(after))
        pPr.append(sp)
        r = p.add_run(txt)
        r.bold = bold; r.font.size = Pt(size)
        r.font.color.rgb = color or DGRAY
        return p

    p0 = cell.paragraphs[0]
    no_space_before(p0)
    pPr = p0._p.get_or_add_pPr()
    sp_el = OxmlElement('w:spacing'); sp_el.set(qn('w:before'), '0'); sp_el.set(qn('w:after'), '60')
    pPr.append(sp_el)
    r0 = p0.add_run(role)
    r0.font.size = Pt(8); r0.font.color.rgb = MUTED

    sp_line(cell, '_' * 42, before=80, after=20)
    sp_line(cell, 'Printed Name & Title', color=MUTED, size=8.5)
    sp_line(cell, ' ', before=60, after=0)
    sp_line(cell, 'Signature: ' + '_' * 30, before=60, after=20)
    sp_line(cell, 'Date: ' + '_' * 34, before=20, after=60)

    cell.width = Cm(8.0)

doc.add_paragraph()

# Footer note
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._p.get_or_add_pPr()
sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '120'); sp.set(qn('w:after'), '0')
pPr.append(sp)
# Top border
pBdr = OxmlElement('w:pBdr')
top = OxmlElement('w:top')
top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '4')
top.set(qn('w:space'), '1'); top.set(qn('w:color'), 'D1D5DB')
pBdr.append(top); pPr.append(pBdr)

r = p.add_run('CAS Enterprise Resource Planning System  ·  Software Proposal  ·  June 16, 2026')
r.font.size = Pt(8.5); r.font.color.rgb = MUTED

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
no_space_before(p2)
r2 = p2.add_run('This document is confidential and intended solely for the named recipient.')
r2.font.size = Pt(8); r2.font.color.rgb = MUTED


# ════════════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════════════

doc.save(OUT_PATH)
print(f'Saved: {OUT_PATH}')
